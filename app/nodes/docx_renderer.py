"""
app/nodes/docx_renderer.py

DOCX tool node — pre-baked variant:
  - DOCXProposalRenderer: render sections to a styled .docx

The .docx twin of PDFProposalRenderer. Identical contract and conventions:
same config (sections / template / proposal_title / client_name), same storage
discipline (workflow state holds the MinIO key, NOT the bytes), same run
signature (state, resolved_config) and same object-store API (put_bytes).

Implementation library: python-docx — same document-generation category as the
python-pptx / openpyxl / WeasyPrint tools already in the stack (a new VARIANT,
not a new tool category).

YAML usage (drop-in swap for PDFProposalRenderer at the end of a workflow):

    - id: generate_docx
      type: DOCXProposalRenderer
      config:
        sections:
          "1. Digital Tool Architecture": "{{tool_architecture.parsed.markdown}}"
          ...
        template: corporate
        proposal_title: "..."
        client_name: "..."

Remember to import this module in app/nodes/__init__.py so the
@NodeRegistry.register decorator fires at import time.
"""
from __future__ import annotations

import asyncio
import io
import re
import uuid
from typing import Any, Literal

from pydantic import BaseModel, Field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

from app.nodes.base import NodeType
from app.nodes.registry import NodeRegistry
from app.observability.logging import get_logger

log = get_logger(__name__)


# --------------------------------------------------------------------------- #
# Template presets — the .docx analogue of the three WeasyPrint CSS templates.
# Same three names as the PDF renderer so `template` is interchangeable.
# --------------------------------------------------------------------------- #
TemplateName = Literal["corporate", "professional", "warm"]

_TEMPLATES: dict[str, dict[str, Any]] = {
    "corporate": {
        "body_font": "Calibri",
        "heading_font": "Calibri Light",
        "accent": RGBColor(0x0D, 0x1B, 0x2A),    # navy
        "accent_2": RGBColor(0xC8, 0xA9, 0x6E),  # gold
        "body_pt": 11,
    },
    "professional": {
        "body_font": "Georgia",
        "heading_font": "Georgia",
        "accent": RGBColor(0x1A, 0x3A, 0x5A),
        "accent_2": RGBColor(0x4A, 0x6A, 0x8A),
        "body_pt": 11,
    },
    "warm": {
        "body_font": "Cambria",
        "heading_font": "Cambria",
        "accent": RGBColor(0x7A, 0x3B, 0x2E),
        "accent_2": RGBColor(0xC0, 0x7A, 0x40),
        "body_pt": 11,
    },
}
_DEFAULT_TEMPLATE = "corporate"


class DOCXRenderInput(BaseModel):
    """Pydantic model defining the DOCXRenderInput shape."""
    pass


class DOCXRenderConfig(BaseModel):
    """Pydantic model defining the DOCXRenderConfig shape.

    Attributes:
        sections (dict[str, str]).
        template (TemplateName).
        proposal_title (str).
        client_name (str).
    """
    sections: dict[str, str] = Field(
        description="Map of section_name -> section_text (Markdown). Templated from upstream nodes."
    )
    template: TemplateName = "corporate"
    proposal_title: str
    client_name: str


class DOCXRenderOutput(BaseModel):
    """Pydantic model defining the DOCXRenderOutput shape.

    Attributes:
        minio_key (str).
        byte_size (int).
        template_used (str).
    """
    minio_key: str
    byte_size: int
    template_used: str


@NodeRegistry.register
class DOCXProposalRenderer(NodeType):
    """Workflow node type implementing the DOCXProposalRenderer capability."""
    type_name = "DOCXProposalRenderer"
    description = "Render proposal sections to a styled .docx (corporate, professional, warm)."
    input_schema = DOCXRenderInput
    output_schema = DOCXRenderOutput
    config_schema = DOCXRenderConfig

    @classmethod
    def required_services(cls, config: dict[str, Any]) -> set[str]:
        """Compute the required services.

        Args:
            config (dict[str, Any]): Node configuration mapping.

        Returns:
            set[str]: The services.
        """
        return {"object_store"}

    async def run(self, state, resolved_config: dict[str, Any]) -> dict[str, Any]:
        """Run the result.

        Args:
            state: Current workflow state.
            resolved_config (dict[str, Any]): Configuration after template resolution.

        Returns:
            dict[str, Any]: The result.
        """
        store = self.services["object_store"]
        cfg = DOCXRenderConfig(**resolved_config)

        # python-docx is sync and CPU-bound — push to a thread, like WeasyPrint.
        docx_bytes = await asyncio.to_thread(
            self._render_docx_bytes,
            cfg.sections,
            cfg.template,
            cfg.client_name,
            cfg.proposal_title,
        )

        # Upload to MinIO under a workflow-scoped key (mirrors the PDF renderer).
        run_id = state.get("inputs", {}).get("SYSTEM.run_id", str(uuid.uuid4()))
        minio_key = f"workflows/{run_id}/proposal.docx"
        await asyncio.to_thread(
            store.put_bytes,
            docx_bytes,
            minio_key,
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        log.info("docx.rendered", node_id=self.node_id,
                 minio_key=minio_key, byte_size=len(docx_bytes),
                 template=cfg.template)
        return {
            "minio_key": minio_key,
            "byte_size": len(docx_bytes),
            "template_used": cfg.template,
        }

    # ----------------------------------------------------------------------- #
    # Document construction (sync; called via asyncio.to_thread)
    # ----------------------------------------------------------------------- #
    def _render_docx_bytes(
        self,
        sections: dict[str, str],
        template_name: str,
        client: str,
        title: str,
    ) -> bytes:
        """Render the docx bytes.

        Args:
            sections (dict[str, str]): The sections.
            template_name (str): The template name.
            client (str): Client instance.
            title (str): The title.

        Returns:
            bytes: The docx bytes.
        """
        tpl = _TEMPLATES.get(template_name, _TEMPLATES[_DEFAULT_TEMPLATE])
        doc = Document()

        # A4 (Horizon Europe submissions are A4), 1-inch margins.
        sec = doc.sections[0]
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(sec, m, Inches(1))

        normal = doc.styles["Normal"]
        normal.font.name = tpl["body_font"]
        normal.font.size = Pt(tpl["body_pt"])

        self._add_title_block(doc, title, client, tpl)

        for heading, body in sections.items():
            h = doc.add_heading(heading, level=1)
            self._style_heading(h, tpl, top_level=True)
            self._render_markdown(doc, body or "", tpl)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_title_block(self, doc, title, client, tpl) -> None:
        """Add the title block.

        Args:
            doc: Document.
            title: The title.
            client: Client instance.
            tpl: The tpl.
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = tpl["heading_font"]
        run.font.color.rgb = tpl["accent"]

        if client:
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = pc.add_run(client)
            rc.font.size = Pt(13)
            rc.font.color.rgb = tpl["accent_2"]

        doc.add_page_break()

    def _style_heading(self, paragraph, tpl, top_level: bool) -> None:
        """Internal helper for the style heading step.

        Args:
            paragraph: The paragraph.
            tpl: The tpl.
            top_level (bool): The top level.
        """
        for run in paragraph.runs:
            run.font.name = tpl["heading_font"]
            run.font.color.rgb = tpl["accent"] if top_level else tpl["accent_2"]

    # ----------------------------------------------------------------------- #
    # Minimal Markdown -> docx renderer. Handles the subset the drafters emit:
    #   # / ## headings, - / * bullets, 1. numbered, **bold** *italic* `code`,
    #   | pipe | tables | (with --- separator), blank-line paragraph breaks.
    # ----------------------------------------------------------------------- #
    _BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
    _NUM_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
    _H_RE = re.compile(r"^(#{1,6})\s+(.*)$")
    _TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}.*$")

    def _render_markdown(self, doc, md: str, tpl) -> None:
        """Render the markdown.

        Args:
            doc: Document.
            md (str): The md.
            tpl: The tpl.
        """
        lines = md.replace("\r\n", "\n").split("\n")
        i, n = 0, len(md.replace("\r\n", "\n").split("\n"))
        while i < n:
            line = lines[i]

            if not line.strip():
                i += 1
                continue

            m = self._H_RE.match(line)
            if m:
                level = min(len(m.group(1)), 4) + 1
                h = doc.add_heading(m.group(2).strip(), level=min(level, 4))
                self._style_heading(h, tpl, top_level=False)
                i += 1
                continue

            if "|" in line and i + 1 < n and self._TABLE_SEP_RE.match(lines[i + 1]):
                i = self._render_table(doc, lines, i, tpl)
                continue

            if self._BULLET_RE.match(line):
                while i < n and self._BULLET_RE.match(lines[i]):
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_inline(p, self._BULLET_RE.match(lines[i]).group(1))
                    i += 1
                continue

            if self._NUM_RE.match(line):
                while i < n and self._NUM_RE.match(lines[i]):
                    p = doc.add_paragraph(style="List Number")
                    self._add_inline(p, self._NUM_RE.match(lines[i]).group(1))
                    i += 1
                continue

            buf = [line]
            i += 1
            while i < n and lines[i].strip() and not (
                self._H_RE.match(lines[i])
                or self._BULLET_RE.match(lines[i])
                or self._NUM_RE.match(lines[i])
                or ("|" in lines[i] and i + 1 < n and self._TABLE_SEP_RE.match(lines[i + 1]))
            ):
                buf.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            self._add_inline(p, " ".join(s.strip() for s in buf))

    def _render_table(self, doc, lines, start, tpl) -> int:
        """Render the table.

        Args:
            doc: Document.
            lines: The lines.
            start: Start value.
            tpl: The tpl.

        Returns:
            int: The table.
        """
        def cells(row: str) -> list[str]:
            """Compute the cells.

            Args:
                row (str): Table row.

            Returns:
                list[str]: The result.
            """
            row = row.strip()
            if row.startswith("|"):
                row = row[1:]
            if row.endswith("|"):
                row = row[:-1]
            return [c.strip() for c in row.split("|")]

        header = cells(lines[start])
        i = start + 2
        body_rows: list[list[str]] = []
        while i < len(lines) and "|" in lines[i] and lines[i].strip():
            body_rows.append(cells(lines[i]))
            i += 1

        cols = len(header)
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"

        hdr = table.rows[0].cells
        for c, text in enumerate(header[:cols]):
            self._add_inline(hdr[c].paragraphs[0], text, bold=True)

        for r in body_rows:
            out = table.add_row().cells
            for c in range(cols):
                self._add_inline(out[c].paragraphs[0], r[c] if c < len(r) else "")

        doc.add_paragraph()
        return i

    _INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")

    def _add_inline(self, paragraph, text: str, bold: bool = False) -> None:
        """Add the inline.

        Args:
            paragraph: The paragraph.
            text (str): The text.
            bold (bool): The bold (optional, default False).
        """
        for token in self._INLINE_RE.split(text):
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                r = paragraph.add_run(token[2:-2]); r.bold = True
            elif token.startswith("`") and token.endswith("`"):
                r = paragraph.add_run(token[1:-1]); r.font.name = "Consolas"
            elif token.startswith("*") and token.endswith("*"):
                r = paragraph.add_run(token[1:-1]); r.italic = True
            else:
                r = paragraph.add_run(token)
            if bold:
                r.bold = True