"""Secure Horizon proposal HTML/Markdown to editable DOCX rendering.

The renderer intentionally shares the HTML sanitisation and PDF-equivalent
pagination pass with :mod:`app.tools.proposal_rendering`.  The Word document is
then built with native OOXML structures: heading styles, numbering, tables,
inline figures, hyperlinks, page fields, and a real TOC field.
"""
from __future__ import annotations

import base64
import hashlib
import io
import re
from dataclasses import dataclass, field
from datetime import date
from typing import Any, Literal

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor
import pdfplumber

from app.tools.proposal_rendering import (
    _sanitise_fragment,
    _toc,
    render_horizon_proposal,
)


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

_EU_BLUE = "003399"
_EU_GOLD = "FFCC00"
_INK = "111827"
_MUTED = "4B5563"
_PALE_BLUE = "EAF0FA"
_LIGHT_BLUE = "F8FAFC"
_BORDER = "9CA3AF"
_WARNING = "9F1239"
_WARNING_BG = "FFF1F2"
_INPUT_ORANGE = "9A3412"
_INPUT_BG = "FFF7ED"

# A4 210 mm, with 20 mm left and 17 mm right margins.
_CONTENT_WIDTH_DXA = 9808
_TABLE_INDENT_DXA = 120
_TABLE_WIDTH_DXA = _CONTENT_WIDTH_DXA - _TABLE_INDENT_DXA
_CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
_BOOKMARK_SAFE = re.compile(r"[^A-Za-z0-9_]+")


@dataclass(frozen=True)
class ProposalDocxRenderResult:
    docx: bytes
    source_html: str
    estimated_page_count: int
    page_count_basis: str
    docx_sha256: str
    source_html_sha256: str
    table_of_contents: list[dict[str, Any]]
    warnings: list[str]
    embedded_image_count: int
    image_placeholder_count: int


@dataclass
class _InlineStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    code: bool = False
    superscript: bool = False
    subscript: bool = False
    input_needed: bool = False
    marked: bool = False


@dataclass
class _TablePlacement:
    row: int
    column: int
    rowspan: int
    colspan: int
    tag: Tag
    header: bool


@dataclass
class _BuildState:
    warnings: list[str] = field(default_factory=list)
    embedded_images: int = 0
    image_placeholders: int = 0
    body_h1_count: int = 0
    bookmark_id: int = 1


def _set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_shading(element, fill: str) -> None:
    properties = (
        element.get_or_add_pPr()
        if element.tag == qn("w:p")
        else element.get_or_add_tcPr()
    )
    existing = properties.find(qn("w:shd"))
    if existing is not None:
        properties.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_run_shading(run, fill: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    existing = r_pr.find(qn("w:shd"))
    if existing is not None:
        r_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    r_pr.append(shading)


def _set_paragraph_borders(
    paragraph,
    *,
    left: tuple[str, int, int] | None = None,
    bottom: tuple[str, int, int] | None = None,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge, spec in (("left", left), ("bottom", bottom)):
        if spec is None:
            continue
        color, size, space = spec
        item = borders.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), str(size))
        item.set(qn("w:space"), str(space))
        item.set(qn("w:color"), color)


def _set_cell_shading(cell, fill: str) -> None:
    _set_shading(cell._tc, fill)


def _set_cell_margins(cell, margins: dict[str, int] | None = None) -> None:
    margins = margins or _CELL_MARGIN_DXA
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[edge]))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _set_table_borders(
    table,
    *,
    style: str = "single",
    color: str = _BORDER,
    size: int = 4,
) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_table_geometry(
    table,
    widths_dxa: list[int],
    *,
    indent_dxa: int = _TABLE_INDENT_DXA,
) -> None:
    """Apply fixed Word table geometry with matching grid and cell widths."""

    if not widths_dxa:
        return
    table.autofit = False
    total_width = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells[: len(widths_dxa)]):
            _set_cell_width(cell, widths_dxa[index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _column_widths(
    rows: list[list[str]],
    column_count: int,
    *,
    total_width: int = _TABLE_WIDTH_DXA,
) -> list[int]:
    if column_count <= 1:
        return [total_width]
    weights: list[int] = []
    for column in range(column_count):
        longest = max(
            (
                len(row[column].strip())
                for row in rows
                if column < len(row)
            ),
            default=12,
        )
        weights.append(max(10, min(70, longest + 4)))

    minimum = min(900, total_width // (column_count * 2))
    remaining = total_width - (minimum * column_count)
    weight_sum = max(1, sum(weights))
    widths = [
        minimum + int(remaining * weight / weight_sum)
        for weight in weights
    ]
    widths[-1] += total_width - sum(widths)
    return widths


def _set_update_fields(doc: DocxDocument) -> None:
    settings = doc.settings._element
    marker = settings.find(qn("w:updateFields"))
    if marker is None:
        marker = OxmlElement("w:updateFields")
        settings.append(marker)
    marker.set(qn("w:val"), "true")


def _add_field(paragraph, instruction: str, fallback: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_node, separate, text, end])


def _add_toc_field(paragraph) -> None:
    _add_field(
        paragraph,
        'TOC \\o "1-3" \\h \\z \\u',
        "Table of contents updates when opened in Word.",
    )


def _add_external_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _EU_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    if bold:
        run_properties.append(OxmlElement("w:b"))
    if italic:
        run_properties.append(OxmlElement("w:i"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _bookmark_name(name: str, bookmark_id: int) -> str:
    safe_name = _BOOKMARK_SAFE.sub("_", name).strip("_")[:38]
    if not safe_name:
        safe_name = f"section_{bookmark_id}"
    if safe_name[0].isdigit():
        safe_name = f"section_{safe_name}"
    return safe_name


def _add_internal_hyperlink(
    paragraph,
    text: str,
    anchor: str,
    *,
    bookmark_id: int,
    bold: bool = False,
    size: float = 10,
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), _bookmark_name(anchor, bookmark_id))
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _INK)
    run_properties.append(color)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_node)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    safe_name = _bookmark_name(name, bookmark_id)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), safe_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _toc_with_pdf_pages(
    entries: list[dict[str, Any]],
    pdf: bytes,
) -> list[dict[str, Any]]:
    """Attach deterministic page numbers from the matching PDF layout."""

    with pdfplumber.open(io.BytesIO(pdf)) as document:
        page_texts = [
            re.sub(
                r"\s+",
                " ",
                (page.extract_text() or "").casefold(),
            )
            for page in document.pages
        ]

    output: list[dict[str, Any]] = []
    for entry in entries:
        title = re.sub(r"\s+", " ", str(entry["title"]).casefold()).strip()
        matching_pages = [
            index
            for index, page_text in enumerate(page_texts, start=1)
            if title and title in page_text
        ]
        output.append(
            {
                **entry,
                # The entry also appears on the TOC page; the last occurrence
                # is the proposal-body heading.
                "page": matching_pages[-1] if matching_pages else None,
            }
        )
    return output


def _set_alt_text(inline_shape, alt_text: str) -> None:
    description = (alt_text or "Proposal figure").strip()[:1000]
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description[:255])


def _normalise_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    return {
        "proposal_title": "Untitled proposal",
        "acronym": "",
        "call_id": "",
        "topic_title": "",
        "action_type": "",
        "consortium_name": "",
        "coordinator": "",
        "version": "Draft",
        "document_date": date.today().isoformat(),
        "part_label": "Part B",
        "confidentiality": "Confidential",
        **metadata,
    }


class _HorizonDocxBuilder:
    def __init__(
        self,
        *,
        metadata: dict[str, Any],
        max_embedded_image_bytes: int,
    ) -> None:
        self.metadata = _normalise_metadata(metadata)
        self.max_embedded_image_bytes = max_embedded_image_bytes
        self.state = _BuildState()
        self.doc = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(18)
        section.right_margin = Mm(17)
        section.bottom_margin = Mm(19)
        section.left_margin = Mm(20)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)
        section.different_first_page_header_footer = True

        self._configure_styles()
        self._configure_header_footer(section)
        _set_update_fields(self.doc)

        properties = self.doc.core_properties
        properties.title = str(self.metadata["proposal_title"])
        properties.subject = "Horizon Europe proposal"
        properties.author = "Eurskem AI"
        properties.keywords = "Horizon Europe, proposal, Part B"
        properties.comments = (
            "Generated from sanitised proposal HTML/Markdown; verify final "
            "pagination in Microsoft Word before submission."
        )

    def _configure_styles(self) -> None:
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.font.color.rgb = RGBColor.from_string(_INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.widow_control = True

        heading_tokens = {
            "Heading 1": (17, _EU_BLUE, 12, 6),
            "Heading 2": (13, _EU_BLUE, 10, 4),
            "Heading 3": (11, "1F2937", 8, 3),
            "Heading 4": (10, _INK, 6, 2),
        }
        for style_name, (size, color, before, after) in heading_tokens.items():
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True

        for style_name in ("List Bullet", "List Number"):
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(10)
            style.paragraph_format.left_indent = Mm(9.5)
            style.paragraph_format.first_line_indent = Mm(-4.9)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.12

        caption = styles["Caption"]
        caption.font.name = "Arial"
        caption.font.size = Pt(8.5)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(_MUTED)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(6)
        caption.paragraph_format.keep_with_next = False

        if "Proposal TOC Heading" not in styles:
            toc_heading = styles.add_style(
                "Proposal TOC Heading",
                WD_STYLE_TYPE.PARAGRAPH,
            )
        else:
            toc_heading = styles["Proposal TOC Heading"]
        toc_heading.font.name = "Arial"
        toc_heading.font.size = Pt(18)
        toc_heading.font.bold = True
        toc_heading.font.color.rgb = RGBColor.from_string(_EU_BLUE)
        toc_heading.paragraph_format.space_before = Pt(0)
        toc_heading.paragraph_format.space_after = Pt(12)

    def _configure_header_footer(self, section) -> None:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        label = (
            f"{self.metadata.get('acronym') or self.metadata['proposal_title']} "
            f"- {self.metadata.get('part_label') or 'Part B'}"
        )
        run = paragraph.add_run(label)
        _set_run_font(run, size=8, color=_MUTED)
        _set_paragraph_borders(
            paragraph,
            bottom=("D1D5DB", 4, 2),
        )

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph._element.getparent().remove(
            footer_paragraph._element
        )
        table = footer.add_table(rows=1, cols=2, width=Mm(173))
        _set_table_geometry(
            table,
            [6200, _TABLE_WIDTH_DXA - 6200],
            indent_dxa=0,
        )
        _set_table_borders(table, style="nil", color="FFFFFF", size=0)
        left = table.cell(0, 0).paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        footer_labels = [
            str(value)
            for value in (
                self.metadata.get("call_id"),
                self.metadata.get("version"),
            )
            if value
        ]
        left_run = left.add_run(" - ".join(footer_labels))
        _set_run_font(left_run, size=8, color=_MUTED)
        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.space_after = Pt(0)
        prefix = right.add_run("Page ")
        _set_run_font(prefix, size=8, color=_MUTED)
        _add_field(right, "PAGE", "1")
        middle = right.add_run(" of ")
        _set_run_font(middle, size=8, color=_MUTED)
        _add_field(right, "NUMPAGES", "1")

        # The cover intentionally has no running furniture.
        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""

    def add_cover(self, evidence_blockers: list[str]) -> None:
        rule = self.doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(0)
        rule.paragraph_format.space_after = Pt(36)
        rule.paragraph_format.line_spacing = 1
        _set_shading(rule._p, _EU_BLUE)
        _set_paragraph_borders(rule, bottom=(_EU_GOLD, 18, 0))
        rule.add_run(" ")

        programme = self.doc.add_paragraph()
        programme.alignment = WD_ALIGN_PARAGRAPH.CENTER
        programme.paragraph_format.space_after = Pt(8)
        run = programme.add_run("HORIZON EUROPE")
        _set_run_font(run, size=12, color=_EU_BLUE, bold=True)

        part = self.doc.add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part.paragraph_format.space_after = Pt(36)
        run = part.add_run(str(self.metadata.get("part_label") or "Part B"))
        _set_run_font(run, size=10, color=_MUTED, bold=True)

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(8)
        run = title.add_run(str(self.metadata["proposal_title"]))
        _set_run_font(run, size=25, color=_EU_BLUE, bold=True)

        acronym = str(self.metadata.get("acronym") or "").strip()
        if acronym:
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(30)
            run = paragraph.add_run(acronym)
            _set_run_font(run, size=17, color=_INK, bold=True)

        rows = [
            ("Call / topic", self.metadata.get("call_id")),
            ("Topic title", self.metadata.get("topic_title")),
            ("Type of action", self.metadata.get("action_type")),
            ("Consortium", self.metadata.get("consortium_name")),
            ("Coordinator", self.metadata.get("coordinator")),
            ("Version", self.metadata.get("version")),
            ("Date", self.metadata.get("document_date")),
        ]
        rows = [(label, value) for label, value in rows if value]
        table = self.doc.add_table(rows=len(rows), cols=2)
        _set_table_geometry(table, [2250, _TABLE_WIDTH_DXA - 2250])
        _set_table_borders(table)
        for index, (label, value) in enumerate(rows):
            label_cell, value_cell = table.rows[index].cells
            _set_cell_shading(label_cell, _PALE_BLUE)
            label_paragraph = label_cell.paragraphs[0]
            label_paragraph.paragraph_format.space_after = Pt(0)
            label_run = label_paragraph.add_run(str(label))
            _set_run_font(label_run, size=9.5, color=_EU_BLUE, bold=True)
            value_paragraph = value_cell.paragraphs[0]
            value_paragraph.paragraph_format.space_after = Pt(0)
            value_run = value_paragraph.add_run(str(value))
            _set_run_font(value_run, size=9.5, color=_INK)

        status = self.doc.add_paragraph()
        status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status.paragraph_format.space_before = Pt(24)
        status.paragraph_format.space_after = Pt(0)
        status_run = status.add_run(
            str(self.metadata.get("confidentiality") or "Confidential")
        )
        _set_run_font(status_run, size=9, color=_MUTED, bold=True)

        if evidence_blockers:
            warning = self.doc.add_paragraph()
            warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
            warning.paragraph_format.space_before = Pt(10)
            warning.paragraph_format.space_after = Pt(0)
            _set_shading(warning._p, _WARNING_BG)
            _set_paragraph_borders(
                warning,
                left=(_WARNING, 10, 6),
            )
            warning_run = warning.add_run(
                "NOT SUBMISSION-READY - unresolved evidence blockers are "
                "listed in the evidence annex."
            )
            _set_run_font(
                warning_run,
                size=9,
                color=_WARNING,
                bold=True,
            )

        self.doc.add_page_break()

    def add_toc(self, entries: list[dict[str, Any]]) -> None:
        heading = self.doc.add_paragraph(style="Proposal TOC Heading")
        heading.add_run("Table of contents")
        _set_paragraph_borders(
            heading,
            bottom=(_EU_BLUE, 12, 4),
        )
        table = self.doc.add_table(rows=len(entries), cols=2)
        _set_table_geometry(table, [8500, _TABLE_WIDTH_DXA - 8500], indent_dxa=0)
        _set_table_borders(table, style="nil", color="FFFFFF", size=0)
        for index, entry in enumerate(entries):
            level = int(entry.get("level") or 1)
            left, right = table.rows[index].cells
            left_paragraph = left.paragraphs[0]
            left_paragraph.paragraph_format.left_indent = Mm(
                max(0, (level - 1) * 5)
            )
            left_paragraph.paragraph_format.space_after = Pt(2)
            _add_internal_hyperlink(
                left_paragraph,
                str(entry["title"]),
                str(entry["id"]),
                bookmark_id=index + 1,
                bold=level == 1,
                size=10 if level == 1 else (9.5 if level == 2 else 9),
            )
            right_paragraph = right.paragraphs[0]
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_paragraph.paragraph_format.space_after = Pt(2)
            run = right_paragraph.add_run(
                str(entry.get("page") or "")
            )
            _set_run_font(
                run,
                size=10 if level == 1 else 9,
                color=_INK,
                bold=level == 1,
            )

        note = self.doc.add_paragraph()
        note.paragraph_format.space_before = Pt(8)
        note.paragraph_format.space_after = Pt(0)
        note_run = note.add_run(
            "Page numbers reflect the matching renderer layout and should be "
            "checked after substantive edits in Word."
        )
        _set_run_font(note_run, size=8, color=_MUTED, italic=True)
        self.doc.add_page_break()

    def add_fragment(self, fragment: str) -> None:
        soup = BeautifulSoup(fragment, "html.parser")
        for child in soup.contents:
            self._render_block(child)

    def _render_block(self, node, *, list_level: int = 0) -> None:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text:
                paragraph = self.doc.add_paragraph()
                paragraph.add_run(text)
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()
        if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = min(int(name[1]), 4)
            paragraph = self.doc.add_paragraph(style=f"Heading {level}")
            if level == 1:
                if self.state.body_h1_count > 0:
                    paragraph.paragraph_format.page_break_before = True
                self.state.body_h1_count += 1
            self._append_inline(paragraph, node)
            heading_id = str(node.get("id") or f"heading-{self.state.bookmark_id}")
            _add_bookmark(
                paragraph,
                heading_id,
                self.state.bookmark_id,
            )
            self.state.bookmark_id += 1
            return

        if name == "p":
            paragraph = self.doc.add_paragraph()
            # Left alignment is materially more stable across Word and
            # LibreOffice than full justification for generated prose. It
            # avoids visibly stretched word spacing after HTML line wrapping.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._append_inline(paragraph, node)
            return

        if name in {"ul", "ol"}:
            self._render_list(node, list_level=list_level)
            return

        if name == "table":
            self._render_table(node)
            return

        if name == "figure":
            self._render_figure(node)
            return

        if name == "img":
            self._render_image(node)
            return

        if name == "blockquote":
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(5)
            paragraph.paragraph_format.right_indent = Mm(2)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(7)
            _set_shading(paragraph._p, _PALE_BLUE)
            _set_paragraph_borders(
                paragraph,
                left=(_EU_BLUE, 18, 6),
            )
            self._append_inline(paragraph, node)
            return

        if name == "pre":
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(4)
            paragraph.paragraph_format.right_indent = Mm(2)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(7)
            _set_shading(paragraph._p, "F3F4F6")
            run = paragraph.add_run(node.get_text("\n", strip=False))
            _set_run_font(run, name="Courier New", size=8, color=_INK)
            return

        if name == "hr":
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            _set_paragraph_borders(
                paragraph,
                bottom=(_BORDER, 6, 1),
            )
            return

        if name == "dl":
            for child in node.find_all(["dt", "dd"], recursive=False):
                paragraph = self.doc.add_paragraph()
                if child.name == "dd":
                    paragraph.paragraph_format.left_indent = Mm(7)
                self._append_inline(
                    paragraph,
                    child,
                    _InlineStyle(bold=child.name == "dt"),
                )
            return

        if name in {
            "div", "section", "main", "article", "details", "summary",
            "tbody", "thead", "tfoot",
        }:
            for child in node.children:
                self._render_block(child, list_level=list_level)
            return

        if name in {"caption", "figcaption"}:
            paragraph = self.doc.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_inline(paragraph, node)
            return

        # Keep unknown but sanitised containers readable.
        text = node.get_text(" ", strip=True)
        if text:
            paragraph = self.doc.add_paragraph()
            self._append_inline(paragraph, node)

    def _append_inline(
        self,
        paragraph,
        container,
        style: _InlineStyle | None = None,
    ) -> None:
        style = style or _InlineStyle()
        children = (
            container.children
            if isinstance(container, Tag)
            else [container]
        )
        for child in children:
            if isinstance(child, NavigableString):
                text = str(child)
                if not text:
                    continue
                run = paragraph.add_run(text)
                _set_run_font(
                    run,
                    name="Courier New" if style.code else "Arial",
                    size=8.5 if style.code else None,
                    color=_INPUT_ORANGE if style.input_needed else None,
                    bold=True if style.bold or style.input_needed else None,
                    italic=True if style.italic else None,
                )
                if style.underline:
                    run.underline = True
                if style.strike:
                    run.font.strike = True
                if style.superscript:
                    run.font.superscript = True
                if style.subscript:
                    run.font.subscript = True
                if style.input_needed:
                    _set_run_shading(run, _INPUT_BG)
                elif style.marked:
                    _set_run_shading(run, "FFF2CC")
                elif style.code:
                    _set_run_shading(run, "F3F4F6")
                continue
            if not isinstance(child, Tag):
                continue
            name = child.name.lower()
            if name == "br":
                paragraph.add_run().add_break()
                continue
            if name == "a":
                href = str(child.get("href") or "").strip()
                label = child.get_text(" ", strip=True) or href
                if href.startswith(("http://", "https://", "mailto:")):
                    _add_external_hyperlink(
                        paragraph,
                        label,
                        href,
                        bold=style.bold,
                        italic=style.italic,
                    )
                else:
                    self._append_inline(paragraph, child, style)
                continue
            if name in {"ul", "ol"}:
                continue
            classes = set(child.get("class") or [])
            next_style = _InlineStyle(
                bold=style.bold or name in {"strong", "b"},
                italic=style.italic or name in {"em", "i"},
                underline=style.underline or name == "u",
                strike=style.strike or name in {"s", "del"},
                code=style.code or name == "code",
                superscript=style.superscript or name == "sup",
                subscript=style.subscript or name == "sub",
                input_needed=(
                    style.input_needed or "input-needed" in classes
                ),
                marked=style.marked or name == "mark",
            )
            self._append_inline(paragraph, child, next_style)

    def _render_list(self, list_tag: Tag, *, list_level: int) -> None:
        ordered = list_tag.name.lower() == "ol"
        for item in list_tag.find_all("li", recursive=False):
            paragraph = self.doc.add_paragraph(
                style="List Number" if ordered else "List Bullet"
            )
            paragraph.paragraph_format.left_indent = Mm(
                9.5 + (list_level * 6)
            )
            paragraph.paragraph_format.first_line_indent = Mm(-4.9)
            self._append_inline(paragraph, item)
            for nested in item.find_all(["ul", "ol"], recursive=False):
                self._render_list(nested, list_level=list_level + 1)

    def _table_placements(
        self,
        table_tag: Tag,
    ) -> tuple[list[_TablePlacement], int, int]:
        rows = table_tag.find_all("tr")
        placements: list[_TablePlacement] = []
        occupied: set[tuple[int, int]] = set()
        max_column = 0
        for row_index, row in enumerate(rows):
            column = 0
            for cell in row.find_all(["th", "td"], recursive=False):
                while (row_index, column) in occupied:
                    column += 1
                try:
                    rowspan = max(1, int(cell.get("rowspan") or 1))
                    colspan = max(1, int(cell.get("colspan") or 1))
                except (TypeError, ValueError):
                    rowspan, colspan = 1, 1
                placement = _TablePlacement(
                    row=row_index,
                    column=column,
                    rowspan=rowspan,
                    colspan=colspan,
                    tag=cell,
                    header=cell.name == "th",
                )
                placements.append(placement)
                for occupied_row in range(row_index, row_index + rowspan):
                    for occupied_column in range(
                        column,
                        column + colspan,
                    ):
                        occupied.add((occupied_row, occupied_column))
                column += colspan
                max_column = max(max_column, column)
        return placements, len(rows), max_column

    def _render_table(self, table_tag: Tag) -> None:
        caption = table_tag.find("caption", recursive=False)
        if caption is not None:
            paragraph = self.doc.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_inline(paragraph, caption)

        placements, row_count, column_count = self._table_placements(table_tag)
        if not row_count or not column_count:
            return
        table = self.doc.add_table(rows=row_count, cols=column_count)
        _set_table_borders(table)

        for placement in placements:
            if placement.rowspan > 1 or placement.colspan > 1:
                try:
                    table.cell(
                        placement.row,
                        placement.column,
                    ).merge(
                        table.cell(
                            placement.row + placement.rowspan - 1,
                            placement.column + placement.colspan - 1,
                        )
                    )
                except (IndexError, ValueError):
                    self.state.warnings.append(
                        "A complex HTML table span could not be reproduced "
                        "exactly in DOCX."
                    )

        plain_rows = [
            [
                table_tag.find_all("tr")[row].find_all(
                    ["th", "td"],
                    recursive=False,
                )[column].get_text(" ", strip=True)
                if column < len(
                    table_tag.find_all("tr")[row].find_all(
                        ["th", "td"],
                        recursive=False,
                    )
                )
                else ""
                for column in range(column_count)
            ]
            for row in range(row_count)
        ]
        widths = _column_widths(plain_rows, column_count)
        _set_table_geometry(table, widths)

        first_row_headers = False
        for placement in placements:
            cell = table.cell(placement.row, placement.column)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            self._append_inline(
                paragraph,
                placement.tag,
                _InlineStyle(bold=placement.header),
            )
            if placement.header:
                _set_cell_shading(cell, _PALE_BLUE)
                first_row_headers = first_row_headers or placement.row == 0
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string("102A56")
            for run in paragraph.runs:
                run.font.size = Pt(8.5)

        if first_row_headers:
            _set_repeat_header(table.rows[0])
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)

    def _placeholder_text(self, figure: Tag) -> str:
        prompt = str(figure.get("data-image-prompt") or "").strip()
        if not prompt:
            prompt = figure.get_text(" ", strip=True)
        prompt = re.sub(
            r"^Image placeholder\s*[-–—:]\s*prompt:\s*",
            "",
            prompt,
            flags=re.IGNORECASE,
        ).strip()
        return prompt or "Visual to be supplied"

    def _render_placeholder(self, prompt: str) -> None:
        self.state.image_placeholders += 1
        table = self.doc.add_table(rows=1, cols=1)
        _set_table_geometry(table, [_TABLE_WIDTH_DXA])
        _set_table_borders(
            table,
            style="dashed",
            color=_EU_BLUE,
            size=10,
        )
        cell = table.cell(0, 0)
        _set_cell_shading(cell, _LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(14)
        run = paragraph.add_run(
            f"Image placeholder - prompt: {prompt}"
        )
        _set_run_font(
            run,
            size=9,
            color="334155",
            italic=True,
        )
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)

    def _decode_data_image(self, source: str) -> bytes:
        header, separator, payload = source.partition(",")
        if not separator or ";base64" not in header.lower():
            raise ValueError("embedded image must use base64 data URI")
        mime = header[5:].split(";", 1)[0].lower()
        if mime not in {"image/png", "image/jpeg", "image/gif"}:
            raise ValueError(f"unsupported embedded image type: {mime}")
        try:
            raw = base64.b64decode(payload, validate=True)
        except Exception as exc:
            raise ValueError("invalid embedded image base64") from exc
        if len(raw) > self.max_embedded_image_bytes:
            raise ValueError(
                "embedded image exceeds "
                f"{self.max_embedded_image_bytes} bytes"
            )
        if not raw:
            raise ValueError("embedded image is empty")
        return raw

    def _render_image(self, image: Tag, caption: Tag | None = None) -> None:
        source = str(image.get("src") or "")
        alt = str(
            image.get("alt")
            or image.get("title")
            or "Proposal figure"
        )
        if not source.startswith("data:image/"):
            self._render_placeholder(alt)
            return
        try:
            raw = self._decode_data_image(source)
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = caption is not None
            run = paragraph.add_run()
            shape = run.add_picture(io.BytesIO(raw), width=Mm(160))
            _set_alt_text(shape, alt)
            self.state.embedded_images += 1
        except Exception as exc:
            self.state.warnings.append(
                f"Embedded image replaced by placeholder: {exc}"
            )
            self._render_placeholder(alt)
            return
        if caption is not None:
            caption_paragraph = self.doc.add_paragraph(style="Caption")
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_inline(caption_paragraph, caption)

    def _render_figure(self, figure: Tag) -> None:
        classes = set(figure.get("class") or [])
        if (
            "image-placeholder" in classes
            or figure.get("data-image-prompt")
        ):
            self._render_placeholder(self._placeholder_text(figure))
            return
        image = figure.find("img")
        caption = figure.find("figcaption")
        if image is not None:
            self._render_image(image, caption)
            return
        if caption is not None:
            paragraph = self.doc.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._append_inline(paragraph, caption)

    def add_bibliography(
        self,
        citation_registry: list[dict[str, Any]],
    ) -> None:
        if not citation_registry:
            return
        heading = self.doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.page_break_before = True
        heading.add_run("References")
        _add_bookmark(
            heading,
            "references",
            self.state.bookmark_id,
        )
        self.state.bookmark_id += 1
        for index, citation in enumerate(citation_registry, start=1):
            number = citation.get("display_number") or index
            text = (
                citation.get("formatted_citation")
                or citation.get("title")
                or citation.get("citation_id")
                or "Unformatted citation"
            )
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(7)
            paragraph.paragraph_format.first_line_indent = Mm(-7)
            paragraph.paragraph_format.space_after = Pt(5)
            prefix = paragraph.add_run(f"[{number}] ")
            _set_run_font(prefix, size=8.5, color=_EU_BLUE, bold=True)
            run = paragraph.add_run(str(text))
            _set_run_font(run, size=8.5, color=_INK)
            citation_id = str(citation.get("citation_id") or "")
            if citation_id:
                _add_bookmark(
                    paragraph,
                    citation_id,
                    self.state.bookmark_id,
                )
                self.state.bookmark_id += 1

    def add_evidence_annex(
        self,
        evidence_qa: dict[str, Any],
        evidence_blockers: list[str],
    ) -> None:
        heading = self.doc.add_paragraph(style="Heading 1")
        heading.paragraph_format.page_break_before = True
        heading.add_run("Evidence integrity annex")
        _add_bookmark(
            heading,
            "evidence_integrity_annex",
            self.state.bookmark_id,
        )
        self.state.bookmark_id += 1

        if evidence_blockers:
            subheading = self.doc.add_paragraph(style="Heading 2")
            subheading.add_run("Blocking issues")
            for blocker in evidence_blockers:
                paragraph = self.doc.add_paragraph(style="List Bullet")
                _set_shading(paragraph._p, _WARNING_BG)
                run = paragraph.add_run(str(blocker))
                _set_run_font(run, size=9, color=_WARNING)

        scalar_items = [
            (key, value)
            for key, value in evidence_qa.items()
            if isinstance(value, (str, int, float, bool))
        ]
        if scalar_items:
            subheading = self.doc.add_paragraph(style="Heading 2")
            subheading.add_run("Quality summary")
            table = self.doc.add_table(rows=len(scalar_items), cols=2)
            _set_table_geometry(table, [3300, _TABLE_WIDTH_DXA - 3300])
            _set_table_borders(table)
            for index, (key, value) in enumerate(scalar_items):
                label_cell, value_cell = table.rows[index].cells
                _set_cell_shading(label_cell, _PALE_BLUE)
                label = label_cell.paragraphs[0]
                label.paragraph_format.space_after = Pt(0)
                label_run = label.add_run(key.replace("_", " ").title())
                _set_run_font(
                    label_run,
                    size=8.5,
                    color=_EU_BLUE,
                    bold=True,
                )
                value_paragraph = value_cell.paragraphs[0]
                value_paragraph.paragraph_format.space_after = Pt(0)
                value_run = value_paragraph.add_run(str(value))
                _set_run_font(value_run, size=8.5, color=_INK)

    def save(self) -> bytes:
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()


def render_horizon_proposal_docx(
    *,
    content: str,
    content_format: Literal["markdown", "html"],
    metadata: dict[str, Any],
    citation_registry: list[dict[str, Any]] | None = None,
    evidence_qa: dict[str, Any] | None = None,
    evidence_blockers: list[str] | None = None,
    include_toc: bool = True,
    include_bibliography: bool = True,
    include_evidence_annex: bool = False,
    page_limit: int | None = 45,
    max_embedded_image_bytes: int = 10_000_000,
) -> ProposalDocxRenderResult:
    """Create an editable Horizon Part B DOCX from sanitised HTML/Markdown.

    DOCX pagination is renderer-dependent.  To preserve the proposal page-limit
    check used by the PDF node, this function also runs the matching
    HTML-to-PDF layout and reports that result as an explicit estimate.
    """

    citation_registry = citation_registry or []
    evidence_qa = evidence_qa or {}
    evidence_blockers = evidence_blockers or []
    fragment = _sanitise_fragment(content, content_format)
    toc = _toc(fragment) if include_toc else []

    pdf_equivalent = render_horizon_proposal(
        content=content,
        content_format=content_format,
        metadata=metadata,
        citation_registry=citation_registry,
        evidence_qa=evidence_qa,
        evidence_blockers=evidence_blockers,
        include_toc=include_toc,
        include_bibliography=include_bibliography,
        include_evidence_annex=include_evidence_annex,
        page_limit=page_limit,
    )
    toc = (
        _toc_with_pdf_pages(toc, pdf_equivalent.pdf)
        if include_toc
        else []
    )

    builder = _HorizonDocxBuilder(
        metadata=metadata,
        max_embedded_image_bytes=max_embedded_image_bytes,
    )
    builder.add_cover(evidence_blockers)
    if include_toc and toc:
        builder.add_toc(toc)
    builder.add_fragment(fragment)
    if include_bibliography:
        builder.add_bibliography(citation_registry)
    if include_evidence_annex:
        builder.add_evidence_annex(evidence_qa, evidence_blockers)
    docx = builder.save()

    warnings = list(dict.fromkeys([
        *pdf_equivalent.warnings,
        *builder.state.warnings,
    ]))
    return ProposalDocxRenderResult(
        docx=docx,
        source_html=pdf_equivalent.html,
        estimated_page_count=pdf_equivalent.page_count,
        page_count_basis="matching_html_pdf_layout",
        docx_sha256=hashlib.sha256(docx).hexdigest(),
        source_html_sha256=hashlib.sha256(
            pdf_equivalent.html.encode("utf-8")
        ).hexdigest(),
        table_of_contents=toc,
        warnings=warnings,
        embedded_image_count=builder.state.embedded_images,
        image_placeholder_count=builder.state.image_placeholders,
    )
