from __future__ import annotations

import base64
import re
from io import BytesIO
from zipfile import ZipFile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from app.nodes.horizon_docx_renderer import HorizonDOCXProposalRenderer
from app.tools.docx_proposal_rendering import (
    DOCX_CONTENT_TYPE,
    render_horizon_proposal_docx,
)


class StubObjectStore:
    def __init__(self):
        self.blobs = {}

    def put_bytes(self, data, key, content_type=None):
        self.blobs[key] = (data, content_type)


def _data_image() -> str:
    buffer = BytesIO()
    Image.new("RGB", (80, 40), color=(0, 51, 153)).save(
        buffer,
        format="PNG",
    )
    return (
        "data:image/png;base64,"
        + base64.b64encode(buffer.getvalue()).decode("ascii")
    )


def test_docx_renderer_builds_native_word_structure_and_sanitises_html():
    result = render_horizon_proposal_docx(
        content=f"""
<h1>1. Excellence</h1>
<h2>1.1 Objectives</h2>
<p>The target remains <strong>[INPUT NEEDED: validated baseline]</strong>.</p>
<ol><li>First measurable objective</li><li>Second measurable objective</li></ol>
<table>
  <caption>Table 1. Validation KPIs</caption>
  <thead><tr><th>KPI</th><th>Target</th></tr></thead>
  <tbody><tr><td>Accuracy</td><td>90%</td></tr></tbody>
</table>
<figure>
  <img src="{_data_image()}" alt="Validated pilot map">
  <figcaption>Figure 1. Pilot regions.</figcaption>
</figure>
<script>alert('x')</script>
<img src="https://example.com/untrusted.png" alt="Map to be supplied">
<h1>2. Impact</h1>
<p>Impact pathway text with <a href="https://doi.org/10.1000/demo">source</a>.</p>
""",
        content_format="html",
        metadata={
            "proposal_title": "AGRO-THRIVE Proposal",
            "acronym": "AGRO-THRIVE",
            "call_id": "HORIZON-CL6-2026-01-CIRCBIO-09",
        },
        citation_registry=[
            {
                "citation_id": "CIT-0001",
                "source_id": "SRC-1",
                "formatted_citation": "Researcher. Verified source. 2025.",
            }
        ],
        evidence_qa={
            "claims_examined": 1,
            "verified_claims": 1,
        },
        evidence_blockers=[],
        include_toc=True,
        include_bibliography=True,
        include_evidence_annex=True,
        page_limit=45,
    )

    assert result.docx.startswith(b"PK")
    assert result.estimated_page_count >= 4
    assert result.page_count_basis == "matching_html_pdf_layout"
    assert result.embedded_image_count == 1
    assert result.image_placeholder_count == 1
    assert "<script" not in result.source_html
    assert "https://example.com/untrusted.png" not in result.source_html
    assert any("INPUT NEEDED" in item for item in result.warnings)

    document = Document(BytesIO(result.docx))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert "1. Excellence" in paragraphs
    assert "2. Impact" in paragraphs
    assert "References" in paragraphs
    assert "Evidence integrity annex" in paragraphs
    assert any(
        paragraph.style.name == "Heading 1"
        and paragraph.text == "1. Excellence"
        for paragraph in document.paragraphs
    )
    assert len(document.tables) >= 4
    assert len(document.inline_shapes) == 1

    with ZipFile(BytesIO(result.docx)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        footer_xml = "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
    assert "w:updateFields" in settings_xml
    assert 'w:hyperlink w:anchor="section_1_excellence"' in document_xml
    assert 'w:hyperlink w:anchor="section_2_impact"' in document_xml
    assert " PAGE " in footer_xml
    assert " NUMPAGES " in footer_xml
    assert "w:tblHeader" in document_xml
    assert 'w:type="fixed"' in document_xml


def test_docx_renderer_builds_real_word_footnotes_when_enabled():
    citation_registry = [
        {
            "display_number": 1,
            "citation_id": "CIT-0001",
            "formatted_citation": "Researcher A. Verified source. 2025.",
        },
        {
            "display_number": 2,
            "citation_id": "CIT-0002",
            "formatted_citation": "Researcher B. Web source. 2024.",
        },
    ]
    result = render_horizon_proposal_docx(
        content=(
            "# 1. Excellence\n\n"
            "Grounded claim from a paper [1] and a web source [2]. "
            "A numbered aside with no matching citation [9] stays literal.\n"
        ),
        content_format="markdown",
        metadata={"proposal_title": "Footnote Test", "acronym": "FN"},
        citation_registry=citation_registry,
        include_toc=False,
        include_bibliography=True,
        page_limit=None,
        enable_footnotes=True,
    )

    assert result.docx.startswith(b"PK")
    assert not any("footnote" in warning.lower() for warning in result.warnings)

    with ZipFile(BytesIO(result.docx)) as archive:
        names = archive.namelist()
        assert "word/footnotes.xml" in names
        content_types = archive.read("[Content_Types].xml").decode("utf-8")
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")

    assert "wordprocessingml.footnotes+xml" in content_types
    assert "relationships/footnotes" in rels

    body_ids = set(re.findall(r'w:footnoteReference w:id="(\d+)"', document_xml))
    part_ids = set(re.findall(r'w:footnote w:id="(-?\d+)"', footnotes_xml))
    assert body_ids == {"1", "2"}
    assert body_ids <= part_ids
    assert {"-1", "0"} <= part_ids
    assert "w:footnoteReference" not in document_xml.replace(
        'w:footnoteReference w:id="1"', ""
    ).replace('w:footnoteReference w:id="2"', "")

    plain_text = re.sub(r"<[^>]+>", "", document_xml)
    assert "[9]" in plain_text  # unmatched marker stays literal, not hijacked

    # Reopening with python-docx must not raise, and the injected part must
    # be rediscoverable via the standard relationship graph.
    reopened = Document(BytesIO(result.docx))
    footnotes_rel = reopened.part.part_related_by(RT.FOOTNOTES)
    assert footnotes_rel.partname == "/word/footnotes.xml"


async def test_horizon_docx_node_stores_docx_and_sanitised_source():
    store = StubObjectStore()
    node = HorizonDOCXProposalRenderer(
        "render_docx",
        {
            "content": (
                "# 1. Excellence\n\n"
                "Grounded proposal text [1].\n\n"
                "# 2. Impact\n\n"
                "A credible impact pathway."
            ),
            "content_format": "markdown",
            "metadata": {
                "proposal_title": "Proposal",
                "acronym": "TEST",
            },
            "citation_registry": [
                {
                    "citation_id": "CIT-0001",
                    "source_id": "SRC-1",
                    "formatted_citation": "Researcher. Verified source. 2025.",
                }
            ],
            "evidence_qa": {
                "claims_examined": 1,
                "verified_claims": 1,
            },
            "evidence_blockers": [],
            "include_evidence_annex": True,
        },
        services={"object_store": store},
    )
    result = await node.run(
        state={"inputs": {"SYSTEM.run_id": "run-123"}},
        resolved_config=node.config.model_dump(),
    )

    assert result["minio_key"] == "workflows/run-123/proposal.docx"
    assert result["docx_key"] == "workflows/run-123/proposal.docx"
    assert (
        result["source_html_key"]
        == "workflows/run-123/proposal.docx-source.html"
    )
    assert result["estimated_page_count"] >= 4
    assert result["page_count"] == result["estimated_page_count"]
    assert result["submission_ready"] is True
    assert store.blobs[result["docx_key"]][0].startswith(b"PK")
    assert store.blobs[result["docx_key"]][1] == DOCX_CONTENT_TYPE
    assert b"Evidence integrity annex" in (
        store.blobs[result["source_html_key"]][0]
    )
